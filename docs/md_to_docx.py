"""Convert CONCEPTS.md to CONCEPTS.docx with proper formatting."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re


def parse_markdown_to_docx(md_path: str, docx_path: str) -> None:
    """Convert markdown file to Word document with formatting."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Skip horizontal rules
        if line.strip() == '---':
            i += 1
            continue

        # Handle headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:].strip(), level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('## '):
            p = doc.add_heading(line[3:].strip(), level=2)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('### '):
            p = doc.add_heading(line[4:].strip(), level=3)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Handle bullet points
        elif line.strip().startswith('* ') or line.strip().startswith('- '):
            text = line.strip()[2:]
            # Handle bold within bullets
            text = handle_inline_formatting(text)
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)

        # Handle numbered lists
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            text = handle_inline_formatting(text)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)

        # Handle code blocks
        elif line.strip().startswith('`') and not line.strip().startswith('```'):
            text = handle_inline_formatting(line.strip())
            p = doc.add_paragraph()
            add_formatted_text(p, text)

        # Handle regular paragraphs
        elif line.strip():
            text = handle_inline_formatting(line.strip())
            p = doc.add_paragraph()
            add_formatted_text(p, text)

        # Empty line
        else:
            if i > 0 and lines[i-1].strip():  # Add space after content
                doc.add_paragraph()

        i += 1

    doc.save(docx_path)
    print(f"Converted {md_path} to {docx_path}")


def handle_inline_formatting(text: str) -> str:
    """Mark inline formatting for later processing."""
    return text


def add_formatted_text(paragraph, text: str) -> None:
    """Add text with inline formatting (bold, italic, code) to paragraph."""
    # Process bold **text**
    parts = re.split(r'(\*\*[^*]+\*\*)', text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            # Handle italic *text* or _text_
            italic_parts = re.split(r'(\*[^*]+\*)', part)
            for ipart in italic_parts:
                if ipart.startswith('*') and ipart.endswith('*') and len(ipart) > 2:
                    run = paragraph.add_run(ipart[1:-1])
                    run.italic = True
                else:
                    paragraph.add_run(ipart)


if __name__ == '__main__':
    md_file = 'D:/job/Research Aston/GenAI Relaibility/Varity v0.1/docs/CONCEPTS.md'
    docx_file = 'D:/job/Research Aston/GenAI Relaibility/Varity v0.1/docs/CONCEPTS.docx'

    parse_markdown_to_docx(md_file, docx_file)
